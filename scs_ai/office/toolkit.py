from __future__ import annotations

from copy import copy
from dataclasses import dataclass
from datetime import datetime, timezone
import json
from pathlib import Path, PurePosixPath, PureWindowsPath
import re
import shutil
from typing import Any
from uuid import uuid4

import openpyxl
from docx import Document
from openpyxl.styles import Font, PatternFill
from openpyxl.utils.cell import range_boundaries


class OfficePathError(ValueError):
    pass


_VERSION_RE = re.compile(r"^(?P<stem>.+)\.v(?P<version>\d{3})$")


@dataclass(frozen=True)
class _Trace:
    trace_id: str
    path: Path


class OfficeToolkit:
    """Contained XLSX/DOCX tooling for SCS AI.

    All paths are workspace-relative. Mutating existing artifacts produce a
    versioned sibling and never silently overwrite the master.
    """

    _SUPPORT_DIRS = ("users", "jobs", "temp", "exports", "backups", "traces")

    def __init__(self, workspace_root: Path | str) -> None:
        self.workspace_root = Path(workspace_root).expanduser().resolve()

    def ensure_workspace(self) -> "OfficeToolkit":
        self.workspace_root.mkdir(parents=True, exist_ok=True)
        for name in self._SUPPORT_DIRS:
            (self.workspace_root / name).mkdir(parents=True, exist_ok=True)
        return self

    def resolve(self, relative_path: str | Path) -> Path:
        raw = str(relative_path)
        if not raw or raw.strip() in {"", "."}:
            raise OfficePathError("workspace-relative path required")
        if PureWindowsPath(raw).is_absolute() or PurePosixPath(raw).is_absolute() or raw.startswith(("//", "\\\\")):
            raise OfficePathError("absolute paths are not allowed")
        normalized = raw.replace("\\", "/")
        raw_parts = normalized.split("/")
        if any(part in {".", ".."} for part in raw_parts):
            raise OfficePathError("path traversal is not allowed")
        parts = PurePosixPath(normalized).parts
        if any(part in {".", ".."} for part in parts):
            raise OfficePathError("path traversal is not allowed")
        candidate = self.workspace_root.joinpath(*parts)
        self._assert_no_link_escape(candidate)
        resolved = candidate.resolve(strict=False)
        try:
            resolved.relative_to(self.workspace_root)
        except ValueError as exc:
            raise OfficePathError("path escapes workspace") from exc
        return resolved

    def workbook_inspect(self, path: str) -> dict[str, Any]:
        op = "workbook.inspect"
        try:
            source = self._require_spreadsheet(path)
            wb = self._open_spreadsheet(source)
            sheets = []
            for ws in wb.worksheets:
                formulas = sum(
                    1 for row in ws.iter_rows(max_row=min(ws.max_row, 2000),
                                              max_col=min(ws.max_column, 200))
                    for cell in row
                    if isinstance(cell.value, str) and cell.value.startswith("=")
                )
                sheets.append({"name": ws.title, "max_row": ws.max_row, "max_column": ws.max_column, "formula_cells": formulas})
            has_vba = bool(getattr(wb, "vba_archive", None) or getattr(wb, "keep_vba", False))
            return self._result(op, source, None, data={"sheets": sheets, "has_vba": has_vba})
        except Exception as exc:
            return self._failure(op, path, exc)

    def workbook_read_range(self, path: str, *, sheet: str, range: str) -> dict[str, Any]:
        op = "workbook.read_range"
        try:
            source = self._require_spreadsheet(path)
            wb = self._open_spreadsheet(source)
            ws = wb[sheet]
            cells = ws[range]
            if not isinstance(cells, tuple):
                cells = ((cells,),)
            elif cells and not isinstance(cells[0], tuple):
                cells = (cells,)
            values = []
            formulas = []
            for row in cells:
                row_values = []
                for cell in row:
                    row_values.append(cell.value)
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        formulas.append({"cell": cell.coordinate, "formula": cell.value, "calculated_value_verified": False})
                values.append(row_values)
            return self._result(op, source, None, data={"values": values, "formula_cells": formulas})
        except Exception as exc:
            return self._failure(op, path, exc)

    def workbook_write_range(self, path: str, *, sheet: str, range: str, values: list[list[Any]]) -> dict[str, Any]:
        op = "workbook.write_range"
        try:
            source = self._require_spreadsheet(path)
            out = self._versioned_sibling(source)
            shutil.copy2(source, out)
            wb = self._open_spreadsheet(out)
            ws = wb[sheet]
            min_col, min_row, max_col, max_row = range_boundaries(range)
            height, width = max_row - min_row + 1, max_col - min_col + 1
            if len(values) != height or any(len(row) != width for row in values):
                raise ValueError("values dimensions must match target range")
            for r_offset, row in enumerate(values):
                for c_offset, value in enumerate(row):
                    ws.cell(row=min_row + r_offset, column=min_col + c_offset, value=value)
            wb.save(out)
            return self._result(op, source, out, changed={"ranges": [f"{sheet}!{range}"], "sections": []})
        except Exception as exc:
            return self._failure(op, path, exc)

    def workbook_set_formula(self, path: str, *, sheet: str, cell: str, formula: str) -> dict[str, Any]:
        op = "workbook.set_formula"
        try:
            if not isinstance(formula, str) or not formula.startswith("="):
                raise ValueError("formula must start with '='")
            source = self._require_spreadsheet(path)
            out = self._versioned_sibling(source)
            shutil.copy2(source, out)
            wb = self._open_spreadsheet(out)
            wb[sheet][cell] = formula
            wb.save(out)
            return self._result(
                op,
                source,
                out,
                data={"formula_written": True, "formula": formula, "calculated_value_verified": False, "recalculation_required": True},
                changed={"ranges": [f"{sheet}!{cell}"], "sections": []},
            )
        except Exception as exc:
            return self._failure(op, path, exc)

    def workbook_format_range(
        self,
        path: str,
        *,
        sheet: str,
        range: str,
        bold: bool | None = None,
        fill: str | None = None,
        number_format: str | None = None,
    ) -> dict[str, Any]:
        op = "workbook.format_range"
        try:
            source = self._require_spreadsheet(path)
            out = self._versioned_sibling(source)
            shutil.copy2(source, out)
            wb = self._open_spreadsheet(out)
            ws = wb[sheet]
            min_col, min_row, max_col, max_row = range_boundaries(range)
            for row in ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                for cell in row:
                    if bold is not None:
                        updated = copy(cell.font)
                        updated.bold = bold
                        cell.font = updated
                    if fill is not None:
                        cell.fill = PatternFill(fill_type="solid", fgColor=fill)
                    if number_format is not None:
                        cell.number_format = number_format
            wb.save(out)
            return self._result(op, source, out, changed={"ranges": [f"{sheet}!{range}"], "sections": []})
        except Exception as exc:
            return self._failure(op, path, exc)

    def workbook_add_sheet(self, path: str, *, sheet: str) -> dict[str, Any]:
        op = "workbook.add_sheet"
        try:
            source = self._require_spreadsheet(path)
            out = self._versioned_sibling(source)
            shutil.copy2(source, out)
            wb = self._open_spreadsheet(out)
            if sheet in wb.sheetnames:
                raise ValueError("sheet already exists")
            wb.create_sheet(sheet)
            wb.save(out)
            return self._result(op, source, out, changed={"ranges": [], "sections": [sheet]})
        except Exception as exc:
            return self._failure(op, path, exc)

    def workbook_export(self, path: str, *, destination: str) -> dict[str, Any]:
        return self._export_copy("workbook.export", path, destination, ".xlsx")

    def document_inspect(self, path: str) -> dict[str, Any]:
        op = "document.inspect"
        try:
            source = self._require_file(path, ".docx")
            doc = Document(source)
            return self._result(op, source, None, data={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)})
        except Exception as exc:
            return self._failure(op, path, exc)

    def document_read(self, path: str) -> dict[str, Any]:
        op = "document.read"
        try:
            source = self._require_file(path, ".docx")
            doc = Document(source)
            chunks = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.append("\t".join(cell.text for cell in row.cells))
            return self._result(op, source, None, data={"text": "\n".join(chunks)})
        except Exception as exc:
            return self._failure(op, path, exc)

    def document_create(self, path: str, *, title: str | None = None, paragraphs: list[str] | None = None) -> dict[str, Any]:
        op = "document.create"
        try:
            destination = self.resolve(path)
            if destination.suffix.lower() != ".docx":
                raise ValueError("document path must end in .docx")
            destination.parent.mkdir(parents=True, exist_ok=True)
            out = self._next_available(destination)
            doc = Document()
            if title:
                doc.add_heading(title, 0)
            for text in paragraphs or []:
                doc.add_paragraph(text)
            doc.save(out)
            return self._result(op, None, out, changed={"ranges": [], "sections": ["document"]})
        except Exception as exc:
            return self._failure(op, path, exc)

    def document_edit(
        self,
        path: str,
        *,
        replacements: dict[str, str] | None = None,
        append_paragraph: str | None = None,
    ) -> dict[str, Any]:
        op = "document.edit"
        try:
            source = self._require_file(path, ".docx")
            out = self._versioned_sibling(source)
            shutil.copy2(source, out)
            doc = Document(out)
            for old, new in (replacements or {}).items():
                self._replace_doc_text(doc, old, new)
            if append_paragraph:
                doc.add_paragraph(append_paragraph)
            doc.save(out)
            return self._result(op, source, out, changed={"ranges": [], "sections": ["document"]})
        except Exception as exc:
            return self._failure(op, path, exc)

    def document_export(self, path: str, *, destination: str) -> dict[str, Any]:
        return self._export_copy("document.export", path, destination, ".docx")

    def schema(self) -> list[dict[str, Any]]:
        names = (
            "workbook.inspect", "workbook.read_range", "workbook.write_range", "workbook.set_formula",
            "workbook.format_range", "workbook.add_sheet", "workbook.export",
            "document.inspect", "document.read", "document.create", "document.edit", "document.export",
        )
        return [{"name": name, "parameters": {"type": "object", "properties": {}}} for name in names]

    def _export_copy(self, op: str, path: str, destination: str, suffix: str) -> dict[str, Any]:
        try:
            source = self._require_file(path, suffix)
            target = self.resolve(destination)
            if target.suffix.lower() != suffix:
                raise ValueError(f"destination must end in {suffix}")
            target.parent.mkdir(parents=True, exist_ok=True)
            out = self._next_available(target)
            shutil.copy2(source, out)
            return self._result(op, source, out, changed={"ranges": [], "sections": ["export"]})
        except Exception as exc:
            return self._failure(op, path, exc)

    def _require_file(self, path: str, suffix: str) -> Path:
        resolved = self.resolve(path)
        if resolved.suffix.lower() != suffix:
            raise ValueError(f"expected {suffix} artifact")
        if not resolved.is_file():
            raise FileNotFoundError(str(resolved))
        return resolved

    def _require_spreadsheet(self, path: str) -> Path:
        resolved = self.resolve(path)
        if resolved.suffix.lower() not in (".xlsx", ".xlsm"):
            raise ValueError("expected .xlsx or .xlsm artifact")
        if not resolved.is_file():
            raise FileNotFoundError(str(resolved))
        return resolved

    def _open_spreadsheet(self, path: Path):
        """Open a workbook; XLSM keeps VBA in memory (never executed)."""
        keep_vba = path.suffix.lower() == ".xlsm"
        workbook = openpyxl.load_workbook(path, data_only=False, keep_vba=keep_vba)
        if keep_vba:
            workbook.keep_vba = True
        return workbook

    def _assert_no_link_escape(self, candidate: Path) -> None:
        current = self.workspace_root
        try:
            relative = candidate.relative_to(self.workspace_root)
        except ValueError as exc:
            raise OfficePathError("path escapes workspace") from exc
        for part in relative.parts:
            current = current / part
            if current.exists() and current.is_symlink():
                resolved = current.resolve()
                try:
                    resolved.relative_to(self.workspace_root)
                except ValueError as exc:
                    raise OfficePathError("symlink/junction escapes workspace") from exc

    @staticmethod
    def _replace_doc_text(doc: Document, old: str, new: str) -> None:
        for paragraph in doc.paragraphs:
            if old in paragraph.text:
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                if old in paragraph.text:
                    paragraph.text = paragraph.text.replace(old, new)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old in paragraph.text:
                            paragraph.text = paragraph.text.replace(old, new)

    def _versioned_sibling(self, source: Path) -> Path:
        base_stem = source.stem
        match = _VERSION_RE.match(base_stem)
        stem = match.group("stem") if match else base_stem
        version = 2
        while True:
            candidate = source.with_name(f"{stem}.v{version:03d}{source.suffix}")
            if not candidate.exists():
                return candidate
            version += 1

    def _next_available(self, target: Path) -> Path:
        if not target.exists():
            return target
        return self._versioned_sibling(target)

    def _result(
        self,
        operation: str,
        input_path: Path | None,
        output_path: Path | None,
        *,
        data: dict[str, Any] | None = None,
        warnings: list[str] | None = None,
        changed: dict[str, list[str]] | None = None,
    ) -> dict[str, Any]:
        trace = self._write_trace(operation, input_path, output_path, True, warnings or [])
        return {
            "success": True,
            "operation": operation,
            "input_artifact": str(input_path) if input_path else None,
            "output_artifact": str(output_path) if output_path else None,
            "warnings": warnings or [],
            "changed": changed or {"ranges": [], "sections": []},
            "data": data or {},
            "trace_id": trace.trace_id,
            "trace_path": str(trace.path),
        }

    def _failure(self, operation: str, input_path: str, exc: Exception) -> dict[str, Any]:
        try:
            resolved = self.resolve(input_path)
        except Exception:
            resolved = None
        warnings = [f"{type(exc).__name__}: {exc}"]
        trace = self._write_trace(operation, resolved, None, False, warnings)
        return {
            "success": False,
            "operation": operation,
            "input_artifact": str(resolved) if resolved else None,
            "output_artifact": None,
            "warnings": warnings,
            "changed": {"ranges": [], "sections": []},
            "data": {},
            "trace_id": trace.trace_id,
            "trace_path": str(trace.path),
        }

    def _write_trace(self, operation: str, input_path: Path | None, output_path: Path | None, success: bool, warnings: list[str]) -> _Trace:
        trace_id = uuid4().hex
        path = self.workspace_root / "traces" / f"{trace_id}.json"
        path.parent.mkdir(parents=True, exist_ok=True)
        payload = {
            "trace_id": trace_id,
            "operation": operation,
            "success": success,
            "input_artifact": str(input_path) if input_path else None,
            "output_artifact": str(output_path) if output_path else None,
            "warnings": warnings,
            "recorded_at": datetime.now(timezone.utc).isoformat(),
        }
        path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        return _Trace(trace_id=trace_id, path=path)
