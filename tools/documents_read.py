from __future__ import annotations

from io import BytesIO
import pymupdf

from tool_sdk import (
    DefendTool,
    ToolContext,
    ToolResult,
    ToolError,
    ToolErrorCode,
    RiskLevel,
    SideEffect,
    DataClassification,
)
from bootstrap_models import (
    DocumentsReadInput,
    DocumentsReadOutput,
    DocumentMediaType,
)
from tools.documents_store import load_raw, load_meta


class DocumentsReadTool(DefendTool[DocumentsReadInput, DocumentsReadOutput]):
    name = "documents.read"
    description = "Extract text or table content from a previously fetched document (PDF, DOCX, XLSX)."
    version = "1.0.0"

    input_model = DocumentsReadInput
    output_model = DocumentsReadOutput

    permissions = frozenset()
    risk_level = RiskLevel.LOW
    side_effect = SideEffect.READ
    idempotent = True
    parallel_safe = True
    timeout_seconds = 45.0
    max_input_classification = DataClassification.PUBLIC
    max_output_classification = DataClassification.PUBLIC

    async def execute(
        self,
        args: DocumentsReadInput,
        context: ToolContext,
    ) -> ToolResult[DocumentsReadOutput]:
        try:
            meta = load_meta(args.document_id)
            raw = load_raw(args.document_id)
            media = DocumentMediaType(meta.get("media_type", "unknown"))
            # Session ownership enforcement
            if meta.get("scope") == "session":
                owner = meta.get("owner_session_id")
                sess = getattr(context, "session_id", None)
                if owner and sess and owner != sess:
                    return ToolResult(
                        ok=False,
                        error=ToolError(
                            code=ToolErrorCode.PERMISSION_DENIED,
                            message="Session document not owned by this session",
                            retryable=False,
                        ),
                    )

            title = meta.get("title")

            if media == DocumentMediaType.PDF:
                content, page_start, page_end, truncated = self._read_pdf(
                    raw, args.page_start, args.page_end, args.max_chars
                )
                return ToolResult(
                    ok=True,
                    data=DocumentsReadOutput(
                        document_id=args.document_id,
                        media_type=media,
                        title=title,
                        content=content,
                        page_start=page_start,
                        page_end=page_end,
                        truncated=truncated,
                        extracted_chars=len(content),
                    ),
                )

            if media == DocumentMediaType.DOCX:
                content, truncated = self._read_docx(raw, args.max_chars)
                return ToolResult(
                    ok=True,
                    data=DocumentsReadOutput(
                        document_id=args.document_id,
                        media_type=media,
                        title=title,
                        content=content,
                        truncated=truncated,
                        extracted_chars=len(content),
                    ),
                )

            if media in {DocumentMediaType.XLSX, DocumentMediaType.XLSM}:
                content, sheet_name, truncated = self._read_xlsx(
                    raw, args.sheet_name, args.max_chars
                )
                return ToolResult(
                    ok=True,
                    data=DocumentsReadOutput(
                        document_id=args.document_id,
                        media_type=media,
                        title=title,
                        content=content,
                        sheet_name=sheet_name,
                        truncated=truncated,
                        extracted_chars=len(content),
                    ),
                )

            if media in {DocumentMediaType.JPEG, DocumentMediaType.PNG}:
                # V1: metadata only — OCR comes later
                from PIL import Image
                img = Image.open(BytesIO(raw))
                info = f"[Image {media.value.upper()}] size={img.size} mode={img.mode}"
                return ToolResult(
                    ok=True,
                    data=DocumentsReadOutput(
                        document_id=args.document_id,
                        media_type=media,
                        title=title,
                        content=info,
                        truncated=False,
                        extracted_chars=len(info),
                    ),
                )

            if media in {DocumentMediaType.TXT, DocumentMediaType.MD, DocumentMediaType.CSV}:
                try:
                    text = raw.decode("utf-8", errors="replace")
                except Exception:
                    text = raw.decode("latin-1", errors="replace")
                max_c = args.max_chars or 12000
                truncated = len(text) > max_c
                content = text[:max_c]
                return ToolResult(
                    ok=True,
                    data=DocumentsReadOutput(
                        document_id=args.document_id,
                        media_type=media,
                        title=title,
                        content=content,
                        page_start=1,
                        page_end=1,
                        truncated=truncated,
                        extracted_chars=len(content),
                    ),
                )

            return ToolResult(
                ok=False,
                error=ToolError(
                    code=ToolErrorCode.INVALID_INPUT,
                    message=f"Cannot extract text from media type: {media.value}",
                    retryable=False,
                ),
            )

        except FileNotFoundError as e:
            return ToolResult(
                ok=False,
                error=ToolError(
                    code=ToolErrorCode.NOT_FOUND,
                    message=str(e),
                    retryable=False,
                ),
            )
        except Exception as e:
            return ToolResult(
                ok=False,
                error=ToolError(
                    code=ToolErrorCode.INTERNAL_ERROR,
                    message=str(e),
                    retryable=False,
                ),
            )

    def _read_pdf(self, raw: bytes, page_start, page_end, max_chars: int):
        import pymupdf
        import pdfplumber

        doc = pymupdf.open(stream=raw, filetype="pdf")
        total = doc.page_count
        start = page_start or 1
        end = page_end or total
        start = max(1, start)
        end = min(total, end)

        parts: list[str] = []
        # Prefer pdfplumber for table-heavy government reports when possible
        try:
            with pdfplumber.open(BytesIO(raw)) as pdf:
                for i in range(start - 1, end):
                    page = pdf.pages[i]
                    text = page.extract_text() or ""
                    tables = page.extract_tables() or []
                    table_text = []
                    for table in tables:
                        for row in table:
                            table_text.append(
                                " | ".join(cell if cell is not None else "" for cell in row)
                            )
                    block = text
                    if table_text:
                        block += "\n" + "\n".join(table_text)
                    parts.append(f"[Page {i+1}]\n{block}")
        except Exception:
            parts = []
            for i in range(start - 1, end):
                parts.append(f"[Page {i+1}]\n{doc[i].get_text('text')}")

        doc.close()
        content = "\n\n".join(parts)
        truncated = False
        if len(content) > max_chars:
            content = content[:max_chars]
            truncated = True
        return content, start, end, truncated

    def _read_docx(self, raw: bytes, max_chars: int):
        from docx import Document
        document = Document(BytesIO(raw))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        content = "\n".join(parts)
        truncated = False
        if len(content) > max_chars:
            content = content[:max_chars]
            truncated = True
        return content, truncated

    def _read_xlsx(self, raw: bytes, sheet_name: str | None, max_chars: int):
        import openpyxl
        wb = openpyxl.load_workbook(BytesIO(raw), read_only=True, data_only=True)
        name = sheet_name or wb.sheetnames[0]
        if name not in wb.sheetnames:
            wb.close()
            raise ValueError(f"Sheet not found: {name}")
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append(" | ".join("" if c is None else str(c) for c in row))
        wb.close()
        content = "\n".join(rows)
        truncated = False
        if len(content) > max_chars:
            content = content[:max_chars]
            truncated = True
        return content, name, truncated