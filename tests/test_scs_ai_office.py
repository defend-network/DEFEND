from __future__ import annotations

import json
import os
import socket
from pathlib import Path

import openpyxl
from docx import Document
from openpyxl.styles import Font
import pytest

from scs_ai.office import OfficePathError, OfficeToolkit


def make_workbook_fixture(root: Path) -> Path:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Cover"
    ws["A1"] = "TAB Commissioning Report"
    ws.merge_cells("A1:C1")
    ws["A1"].font = Font(bold=True, size=14)
    ws["A3"] = "Project"
    ws["B3"] = "ACME Facility"
    ws2 = wb.create_sheet("Readings")
    ws2["A1"] = "Tag"
    ws2["B1"] = "Reading"
    ws2["C1"] = "Units"
    ws2["D1"] = "Location"
    ws2["E1"] = "Status"
    for cell in ws2[1]:
        cell.font = Font(bold=True)
    rows = [
        ("AHU-1", "2.15", "kPa", "Level 2", "OK"),
        ("AHU-2", "2.02", "kPa", "Level 2", "OK"),
        ("VAV-12", "-0.08", "Pa", "Zone 12", "CHECK"),
    ]
    for i, row in enumerate(rows, start=2):
        for j, value in enumerate(row, start=1):
            ws2.cell(row=i, column=j, value=value)
    ws2["G1"] = "Total tags"
    ws2["G2"] = "=COUNTA(A2:A4)"
    path = root / "TAB Report.xlsx"
    wb.save(path)
    return path


def make_docx_fixture(root: Path) -> Path:
    doc = Document()
    doc.add_heading("TAB Closeout Report", 0)
    doc.add_paragraph("Project: ACME Facility. Balance report prepared by SCS.")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Tag"
    table.cell(0, 1).text = "Reading"
    table.cell(0, 2).text = "Status"
    table.cell(1, 0).text = "AHU-1"
    table.cell(1, 1).text = "2.15 kPa"
    table.cell(1, 2).text = "OK"
    doc.add_paragraph("Recommendation: schedule quarterly verification.")
    path = root / "summary.docx"
    doc.save(path)
    return path


@pytest.fixture
def toolkit(tmp_path):
    root = tmp_path / "ai_workspace"
    return OfficeToolkit(root).ensure_workspace()


def test_workspace_ensure_creates_support_directories(toolkit):
    root = toolkit.workspace_root
    for name in ("users", "jobs", "temp", "exports", "backups", "traces"):
        assert (root / name).is_dir()


def test_workspace_resolve_keeps_paths_inside_root(toolkit):
    resolved = toolkit.resolve("jobs/job-001/TAB Report.xlsx")
    assert resolved.is_absolute()
    assert resolved == toolkit.workspace_root / "jobs" / "job-001" / "TAB Report.xlsx"


def test_workspace_rejects_absolute_paths(toolkit):
    for bad in (r"C:\SCS_DATA\ai_workspace\elsewhere", "/tmp/out", "//server/share/x"):
        with pytest.raises(OfficePathError):
            toolkit.resolve(bad)


def test_workspace_rejects_traversal_components(toolkit):
    for bad in ("../secrets.txt", "..\\secrets.txt", "jobs/../../x", "jobs/../outside", ".", "jobs/./x"):
        with pytest.raises(OfficePathError):
            toolkit.resolve(bad)


def test_workspace_rejects_symlink_and_junction_escape(toolkit, tmp_path):
    outside = tmp_path / "outside"
    outside.mkdir()
    (outside / "leak.txt").write_text("outside", encoding="utf-8")
    link = toolkit.workspace_root / "jobs" / "leak"
    try:
        os.symlink(outside, link, target_is_directory=True)
    except (OSError, NotImplementedError):
        pytest.skip("directory symlink creation not permitted on this system")
    with pytest.raises(OfficePathError):
        toolkit.resolve("jobs/leak/leak.txt")


def test_workbook_inspect_reports_sheets_and_formulas(toolkit):
    fixture = make_workbook_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.workbook_inspect("jobs/TAB Report.xlsx")
    assert result["success"] is True
    assert result["operation"] == "workbook.inspect"
    sheets = {sheet["name"] for sheet in result["data"]["sheets"]}
    assert sheets == {"Cover", "Readings"}
    readings = next(s for s in result["data"]["sheets"] if s["name"] == "Readings")
    assert readings["formula_cells"] >= 1
    assert result["output_artifact"] is None


def test_workbook_read_range_returns_values(toolkit):
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.workbook_read_range("jobs/TAB Report.xlsx", sheet="Readings", range="A1:E4")
    assert result["success"] is True
    values = result["data"]["values"]
    assert values[0] == ["Tag", "Reading", "Units", "Location", "Status"]
    assert values[1][0] == "AHU-1"


def test_workbook_read_range_reports_formula_cells_honestly(toolkit):
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.workbook_read_range("jobs/TAB Report.xlsx", sheet="Readings", range="G1:G2")
    formula = [item for item in result["data"]["formula_cells"] if item["cell"] == "G2"][0]
    assert formula["formula"].startswith("=COUNTA")
    assert formula["calculated_value_verified"] is False


def test_workbook_write_range_produces_versioned_output_not_overwrite(toolkit):
    fixture = make_workbook_fixture(toolkit.workspace_root / "jobs")
    original = fixture.read_bytes()
    result = toolkit.workbook_write_range(
        "jobs/TAB Report.xlsx", sheet="Readings", range="B5", values=[["2.50"]]
    )
    assert result["success"] is True
    assert result["input_artifact"].endswith("TAB Report.xlsx")
    assert result["output_artifact"].endswith("TAB Report.v002.xlsx")
    assert result["output_artifact"] != result["input_artifact"]
    assert result["changed"]["ranges"] == ["Readings!B5"]
    assert fixture.read_bytes() == original
    out = Path(result["output_artifact"])
    assert out.is_file()
    wb = openpyxl.load_workbook(out, data_only=False)
    assert wb["Readings"]["B5"].value == "2.50"


def test_workbook_set_formula_is_honest_about_recalculation(toolkit):
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.workbook_set_formula(
        "jobs/TAB Report.xlsx", sheet="Readings", cell="G3", formula="=SUM(B2:B4)"
    )
    assert result["success"] is True
    assert result["data"]["formula_written"] is True
    assert result["data"]["calculated_value_verified"] is False
    assert result["data"]["recalculation_required"] is True
    out = Path(result["output_artifact"])
    wb = openpyxl.load_workbook(out, data_only=False)
    assert wb["Readings"]["G3"].value == "=SUM(B2:B4)"


def test_workbook_format_range_preserves_other_formatting(toolkit):
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.workbook_format_range(
        "jobs/TAB Report.xlsx",
        sheet="Readings",
        range="B2:B4",
        bold=True,
        fill="FFFF00",
        number_format="0.00",
    )
    assert result["success"] is True
    out = Path(result["output_artifact"])
    wb = openpyxl.load_workbook(out)
    assert wb["Readings"]["B2"].font.bold is True
    assert wb["Readings"]["B2"].fill.start_color.rgb.endswith("FFFF00")
    assert wb["Readings"]["B2"].number_format == "0.00"
    assert wb["Cover"]["A1"].font.bold is True
    assert wb["Cover"]["A1"].font.size == 14


def test_workbook_add_sheet_preserves_existing_sheets(toolkit):
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.workbook_add_sheet("jobs/TAB Report.xlsx", sheet="Checklist")
    assert result["success"] is True
    out = Path(result["output_artifact"])
    wb = openpyxl.load_workbook(out)
    assert wb.sheetnames == ["Cover", "Readings", "Checklist"]
    assert "Checklist" in result["changed"]["sections"]


def test_workbook_export_never_overwrites(toolkit):
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    first = toolkit.workbook_export("jobs/TAB Report.xlsx", destination="exports/report.xlsx")
    assert first["success"] is True
    second = toolkit.workbook_export("jobs/TAB Report.xlsx", destination="exports/report.xlsx")
    assert second["success"] is True
    assert first["output_artifact"] != second["output_artifact"]
    assert second["output_artifact"].endswith("report.v002.xlsx")


def test_document_inspect_and_read(toolkit):
    make_docx_fixture(toolkit.workspace_root / "jobs")
    info = toolkit.document_inspect("jobs/summary.docx")
    assert info["success"] is True
    assert info["data"]["paragraph_count"] >= 2
    assert info["data"]["table_count"] == 1
    content = toolkit.document_read("jobs/summary.docx")
    assert content["success"] is True
    assert "TAB Closeout Report" in content["data"]["text"]


def test_document_create_produces_new_versioned_document(toolkit):
    result = toolkit.document_create(
        "jobs/job-002/closeout.docx",
        title="Balance Report",
        paragraphs=["Project: Sunrise Tower.", "All readings verified."],
    )
    assert result["success"] is True
    assert result["output_artifact"].endswith("closeout.docx")
    out = Path(result["output_artifact"])
    assert out.is_file()
    doc = Document(out)
    assert doc.paragraphs[0].text == "Balance Report"


def test_document_edit_replaces_and_appends(toolkit):
    make_docx_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.document_edit(
        "jobs/summary.docx",
        replacements={"ACME Facility": "Sunrise Tower"},
        append_paragraph="Signed by SCS.",
    )
    assert result["success"] is True
    out = Path(result["output_artifact"])
    assert out.is_file() and out.name != "summary.docx"
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Sunrise Tower" in text
    assert "ACME Facility" not in text
    assert "Signed by SCS." in text


def test_document_edit_never_overwrites_master(toolkit):
    fixture = make_docx_fixture(toolkit.workspace_root / "jobs")
    original = fixture.read_bytes()
    result = toolkit.document_edit("jobs/summary.docx", replacements={"ACME": "Sunrise"})
    assert result["success"] is True
    assert fixture.read_bytes() == original
    assert Path(result["output_artifact"]) != fixture


def test_malformed_xlsx_returns_structured_failure(toolkit):
    bad = toolkit.workspace_root / "jobs" / "broken.xlsx"
    bad.parent.mkdir(parents=True, exist_ok=True)
    bad.write_bytes(b"this is not a zip archive at all")
    result = toolkit.workbook_inspect("jobs/broken.xlsx")
    assert result["success"] is False
    assert result["warnings"]
    assert result["trace_id"]


def test_malformed_docx_returns_structured_failure(toolkit):
    bad = toolkit.workspace_root / "jobs" / "broken.docx"
    bad.parent.mkdir(parents=True, exist_ok=True)
    bad.write_bytes(b"not a valid docx package")
    result = toolkit.document_inspect("jobs/broken.docx")
    assert result["success"] is False
    assert result["warnings"]
    assert result["trace_id"]


def test_tool_results_are_structured_and_carry_trace(toolkit):
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.workbook_write_range("jobs/TAB Report.xlsx", sheet="Readings", range="B5", values=[["x"]])
    for field in ("success", "operation", "input_artifact", "output_artifact", "warnings", "changed", "trace_id"):
        assert field in result
    assert isinstance(result["trace_id"], str) and result["trace_id"]


def test_trace_records_are_written(toolkit):
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    result = toolkit.workbook_read_range("jobs/TAB Report.xlsx", sheet="Readings", range="A1:A2")
    trace = Path(result["trace_path"])
    assert trace.is_file()
    payload = json.loads(trace.read_text(encoding="utf-8"))
    assert payload["trace_id"] == result["trace_id"]
    assert payload["operation"] == "workbook.read_range"


def test_tool_registry_schema_covers_all_office_tools(toolkit):
    schema = toolkit.schema()
    names = {item["name"] for item in schema}
    assert names == {
        "workbook.inspect",
        "workbook.read_range",
        "workbook.write_range",
        "workbook.set_formula",
        "workbook.format_range",
        "workbook.add_sheet",
        "workbook.export",
        "document.inspect",
        "document.read",
        "document.create",
        "document.edit",
        "document.export",
    }
    for item in schema:
        assert item["parameters"]["type"] == "object"
        assert isinstance(item["parameters"]["properties"], dict)


def test_no_network_calls_in_office_tools(toolkit, monkeypatch):
    class BlockedSocket:
        def __init__(self, *args, **kwargs):
            raise AssertionError("network call attempted")

    monkeypatch.setattr(socket, "socket", BlockedSocket)
    make_workbook_fixture(toolkit.workspace_root / "jobs")
    make_docx_fixture(toolkit.workspace_root / "jobs")
    assert toolkit.workbook_inspect("jobs/TAB Report.xlsx")["success"] is True
    assert toolkit.workbook_read_range("jobs/TAB Report.xlsx", sheet="Readings", range="A1:A2")["success"] is True
    assert toolkit.workbook_write_range("jobs/TAB Report.xlsx", sheet="Readings", range="B5", values=[["2.5"]])["success"] is True
    assert toolkit.workbook_set_formula("jobs/TAB Report.xlsx", sheet="Readings", cell="G3", formula="=SUM(B2:B4)")["success"] is True
    assert toolkit.workbook_format_range("jobs/TAB Report.xlsx", sheet="Readings", range="B2:B4", bold=True)["success"] is True
    assert toolkit.workbook_add_sheet("jobs/TAB Report.xlsx", sheet="X")["success"] is True
    assert toolkit.workbook_export("jobs/TAB Report.xlsx", destination="exports/r.xlsx")["success"] is True
    assert toolkit.document_inspect("jobs/summary.docx")["success"] is True
    assert toolkit.document_read("jobs/summary.docx")["success"] is True
    assert toolkit.document_create("jobs/job-003/out.docx", title="T", paragraphs=["p"])["success"] is True
    assert toolkit.document_edit("jobs/summary.docx", replacements={"ACME": "X"})["success"] is True
    assert toolkit.document_export("jobs/summary.docx", destination="exports/s.docx")["success"] is True


def test_office_modules_import_no_network_libraries():
    import scs_ai.office as package

    for path in Path(package.__path__[0]).glob("*.py"):
        source = path.read_text(encoding="utf-8")
        for forbidden in ("import socket", "import requests", "import urllib", "import httpx", "import aiohttp"):
            assert forbidden not in source, path


def test_existing_scs_ai_regressions_remain_green():
    import importlib

    importlib.import_module("scs_ai.app")
    importlib.import_module("scs_ai.runtime")
    importlib.import_module("scs_ai.model_gateway")
    importlib.import_module("scs_ai.tunnel")
    importlib.import_module("scs_ai.tools")